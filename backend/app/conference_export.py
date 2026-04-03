from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def export_docx(
    out_path: Path,
    original: str,
    translated: str,
    src_label: str,
    tgt_label: str,
) -> None:
    doc = Document()
    doc.add_heading("Оригинал", level=1)
    doc.add_paragraph(original or "")
    doc.add_heading("Перевод", level=1)
    doc.add_paragraph(translated or "")
    doc.save(str(out_path))


def export_pdf(
    out_path: Path,
    original: str,
    translated: str,
    src_label: str,
    tgt_label: str,
) -> None:
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))

    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4

    x = 40
    y = height - 40

    def new_page() -> None:
        nonlocal y
        c.showPage()
        c.setFont("DejaVuSans", 11)
        y = height - 40

    def draw_title(title: str) -> None:
        nonlocal y
        c.setFont("DejaVuSans", 14)
        c.drawString(x, y, title)
        y -= 22
        c.setFont("DejaVuSans", 11)

    def draw_text(text: str) -> None:
        nonlocal y
        for line in (text or "").splitlines() or [""]:
            while len(line) > 95:
                c.drawString(x, y, line[:95])
                line = line[95:]
                y -= 14
                if y < 60:
                    new_page()
            c.drawString(x, y, line)
            y -= 14
            if y < 60:
                new_page()
        c.setFont("DejaVuSans", 11)

    draw_title("Оригинал")
    draw_text(original)

    y -= 10
    if y < 80:
        new_page()

    draw_title("Перевод")
    draw_text(translated)

    c.save()
