# backend/app/uploads_export.py
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def export_docx(job, path: Path) -> None:
    doc = Document()
    doc.add_heading(job.filename or "Перевод", 0)

    doc.add_paragraph(f"Исходный язык: {job.src_language or 'auto'}")
    doc.add_paragraph(f"Язык перевода: {job.target_language}")
    doc.add_paragraph("")

    doc.add_heading("Оригинальный текст", level=1)
    doc.add_paragraph(job.transcript_text or "")

    doc.add_paragraph("")
    doc.add_heading("Перевод", level=1)
    doc.add_paragraph(job.translated_text or "")

    doc.save(str(path))


def _export_pdf(
    out_path: Path, src_text: str, trg_text: str, src_label: str, trg_label: str
) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))

    c = canvas.Canvas(str(out_path), pagesize=A4)
    _, height = A4
    x = 40
    y = height - 40

    def new_page():
        nonlocal y
        c.showPage()
        y = height - 40

    def draw_title(title: str):
        nonlocal y
        c.setFont("DejaVuSans", 12)
        c.drawString(x, y, title)
        y -= 18

    def draw_text(text: str):
        nonlocal y
        c.setFont("DejaVuSans", 10)
        for line in (text or "").splitlines() or [""]:
            # перенос строк “по символам” (просто и стабильно)
            while len(line) > 95:
                c.drawString(x, y, line[:95])
                line = line[95:]
                y -= 14
                if y < 60:
                    new_page()
                    c.setFont("DejaVuSans", 10)
            c.drawString(x, y, line)
            y -= 14
            if y < 60:
                new_page()
                c.setFont("DejaVuSans", 10)

    draw_title(f"Оригинал ({src_label})")
    draw_text(src_text)
    y -= 10
    if y < 80:
        new_page()

    draw_title(f"Перевод ({trg_label})")
    draw_text(trg_text)

    c.save()
