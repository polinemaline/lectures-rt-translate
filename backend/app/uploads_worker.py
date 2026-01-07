# backend/app/uploads_worker.py
import os
import re
import subprocess
import threading
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from sqlalchemy.orm import Session

from app.db import SessionLocal
from app.models import UploadJob

UPLOADS_ROOT = Path("/data/uploads")

# ---------------------------
# ASR (Whisper via Transformers)
# ---------------------------
ASR_MODEL = os.getenv("ASR_MODEL", "openai/whisper-small")  # whisper-tiny/base/small
ASR_CHUNK_LENGTH_S = int(os.getenv("ASR_CHUNK_LENGTH_S", "30"))
ASR_STRIDE_LENGTH_S = int(os.getenv("ASR_STRIDE_LENGTH_S", "5"))

# ---------------------------
# Translation (LOCAL: OPUS-MT)
# ---------------------------
DEFAULT_SRC_LANG = os.getenv("DEFAULT_SRC_LANG", "rus_Cyrl")

# Переводим по предложениям: это главный фикс "обрывается перевод"
TR_BATCH_SIZE = int(os.getenv("TR_BATCH_SIZE", "8"))
TR_MAX_NEW_TOKENS = int(
    os.getenv("TR_MAX_NEW_TOKENS", "256")
)  # на предложение обычно хватает
TR_SENT_MAX_CHARS = int(
    os.getenv("TR_SENT_MAX_CHARS", "450")
)  # длинные предложения режем

# ---------------------------
# Caches
# ---------------------------
_ASR_PIPE = None
_TRANSLATORS: Dict[Tuple[str, str], object] = {}  # (src_iso, tgt_iso) -> pipeline


# ---------------------------
# DB helpers (sync)
# ---------------------------
def _db_update(job_id: int, **fields) -> None:
    with SessionLocal() as db:  # type: Session
        job = db.query(UploadJob).filter(UploadJob.id == job_id).first()
        if not job:
            return
        for k, v in fields.items():
            setattr(job, k, v)
        db.commit()


def _db_get_job(job_id: int) -> Optional[UploadJob]:
    with SessionLocal() as db:
        return db.query(UploadJob).filter(UploadJob.id == job_id).first()


# ---------------------------
# File helpers
# ---------------------------
def _find_input_file(base: Path) -> Path:
    for cand in sorted(base.glob("input.*")):
        if cand.is_file():
            return cand
    raise FileNotFoundError("Input file not found (expected input.*)")


def _ffmpeg_extract(input_path: Path, wav_path: Path) -> tuple[int, int]:
    """
    Делает wav 16kHz mono.
    Возвращает: (size_bytes, duration_sec)
    """
    cmd = [
        "ffmpeg",
        "-y",
        "-nostdin",
        "-hide_banner",
        "-loglevel",
        "error",
        "-i",
        str(input_path),
        "-ac",
        "1",
        "-ar",
        "16000",
        "-vn",
        str(wav_path),
    ]
    p = subprocess.run(cmd, capture_output=True, text=True)
    if p.returncode != 0:
        raise RuntimeError(p.stderr.strip() or f"ffmpeg failed (code={p.returncode})")

    size = wav_path.stat().st_size if wav_path.exists() else 0

    dur_cmd = [
        "ffprobe",
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        str(wav_path),
    ]
    d = subprocess.run(dur_cmd, capture_output=True, text=True)
    duration = 0
    if d.returncode == 0:
        try:
            duration = int(float(d.stdout.strip()))
        except Exception:
            duration = 0

    return size, duration


# ---------------------------
# Progress ticker
# ---------------------------
def _progress_ticker(
    job_id: int, stage: str, start: int, end: int, step: int = 1, every_sec: int = 3
):
    stop = threading.Event()

    def run():
        p = start
        while not stop.is_set() and p < end:
            _db_update(job_id, stage=stage, progress=p)
            p += step
            stop.wait(every_sec)

    t = threading.Thread(target=run, daemon=True)
    t.start()
    return stop


# ---------------------------
# ASR
# ---------------------------
def _get_asr_pipeline():
    global _ASR_PIPE
    if _ASR_PIPE is None:
        from transformers import pipeline

        _ASR_PIPE = pipeline(
            "automatic-speech-recognition",
            model=ASR_MODEL,
            device=-1,  # CPU
        )
    return _ASR_PIPE


def _nllb_to_iso(nllb_code: str) -> Optional[str]:
    mapping = {
        "rus_Cyrl": "ru",
        "eng_Latn": "en",
        "deu_Latn": "de",
        "fra_Latn": "fr",
        "spa_Latn": "es",
        "ita_Latn": "it",
        "por_Latn": "pt",
        "tur_Latn": "tr",
    }
    return mapping.get(nllb_code)


def _transcribe_whisper(wav_path: Path, src_lang_nllb: str) -> str:
    asr = _get_asr_pipeline()
    iso = _nllb_to_iso(src_lang_nllb)

    generate_kwargs = {"task": "transcribe"}
    if iso:
        generate_kwargs["language"] = iso

    out = asr(
        str(wav_path),
        chunk_length_s=ASR_CHUNK_LENGTH_S,
        stride_length_s=ASR_STRIDE_LENGTH_S,
        generate_kwargs=generate_kwargs,
        return_timestamps=False,
    )

    return (out.get("text") or "").strip()


# ---------------------------
# Translation (LOCAL OPUS-MT) — FIX: перевод по предложениям
# ---------------------------
def _get_translator(src_iso: str, tgt_iso: str):
    key = (src_iso, tgt_iso)
    if key in _TRANSLATORS:
        return _TRANSLATORS[key]

    from transformers import pipeline

    model_name = f"Helsinki-NLP/opus-mt-{src_iso}-{tgt_iso}"
    tr = pipeline("translation", model=model_name, device=-1)
    _TRANSLATORS[key] = tr
    return tr


def _split_long_string(s: str, max_chars: int) -> List[str]:
    s = s.strip()
    if not s:
        return []
    if len(s) <= max_chars:
        return [s]
    parts: List[str] = []
    cur: List[str] = []
    cur_len = 0
    for w in s.split():
        add = len(w) + (1 if cur else 0)
        if cur_len + add <= max_chars:
            cur.append(w)
            cur_len += add
        else:
            parts.append(" ".join(cur))
            cur = [w]
            cur_len = len(w)
    if cur:
        parts.append(" ".join(cur))
    return parts


_SENT_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+|\n+")


def _sentencize(text: str) -> List[str]:
    """
    Делим на предложения/строки так, чтобы не давать модели огромный кусок.
    """
    text = (text or "").strip()
    if not text:
        return []
    raw = [p.strip() for p in _SENT_SPLIT_RE.split(text) if p.strip()]
    out: List[str] = []
    for p in raw:
        out.extend(_split_long_string(p, TR_SENT_MAX_CHARS))
    return out


def _translate_batch(tr, texts: List[str]) -> List[str]:
    outs = tr(
        texts,
        truncation=True,
        max_new_tokens=TR_MAX_NEW_TOKENS,
        batch_size=TR_BATCH_SIZE,
    )

    res: List[str] = []
    for item in outs:
        if isinstance(item, list) and item and isinstance(item[0], dict):
            res.append((item[0].get("translation_text") or "").strip())
        elif isinstance(item, dict):
            res.append((item.get("translation_text") or "").strip())
        else:
            res.append("")
    return res


def _translate_safe(tr, sentence: str) -> str:
    """
    Страховка: если модель вернула пусто — возвращаем исходник,
    чтобы ничего не “терялось”.
    """
    sentence = (sentence or "").strip()
    if not sentence:
        return ""
    out = _translate_batch(tr, [sentence])[0].strip()
    return out if out else sentence


def _translate_text_opus(text: str, src_iso: str, tgt_iso: str, job_id: int) -> str:
    if not text.strip():
        return ""
    if src_iso == tgt_iso:
        return text

    # 1) пытаемся прямую пару
    try:
        tr = _get_translator(src_iso, tgt_iso)
        sents = _sentencize(text)
        if not sents:
            return ""

        out_parts: List[str] = []
        n = len(sents)

        # батчами
        i = 0
        while i < n:
            batch = sents[i : i + TR_BATCH_SIZE]
            prog = 70 + int(((i + len(batch)) / n) * 18)  # 70..88
            _db_update(job_id, stage="translate", progress=prog)

            translated = _translate_batch(tr, batch)
            # страховка от пустых кусочков
            for src_sent, trg_sent in zip(batch, translated):
                out_parts.append(trg_sent.strip() if trg_sent.strip() else src_sent)

            i += TR_BATCH_SIZE

        return " ".join([p for p in out_parts if p]).strip()

    except Exception:
        # 2) pivot через en
        if src_iso != "en" and tgt_iso != "en":
            tr1 = _get_translator(src_iso, "en")
            tr2 = _get_translator("en", tgt_iso)

            sents = _sentencize(text)
            if not sents:
                return ""

            mid_parts: List[str] = []
            n = len(sents)

            i = 0
            while i < n:
                batch = sents[i : i + TR_BATCH_SIZE]
                prog = 70 + int(((i + len(batch)) / n) * 9)  # 70..79
                _db_update(job_id, stage="translate", progress=prog)

                mid = _translate_batch(tr1, batch)
                for src_sent, mid_sent in zip(batch, mid):
                    mid_parts.append(mid_sent.strip() if mid_sent.strip() else src_sent)
                i += TR_BATCH_SIZE

            mid_text = " ".join([p for p in mid_parts if p]).strip()
            mid_sents = _sentencize(mid_text)
            out_parts: List[str] = []
            n2 = len(mid_sents)

            i = 0
            while i < n2:
                batch = mid_sents[i : i + TR_BATCH_SIZE]
                prog = 80 + int(((i + len(batch)) / n2) * 8)  # 80..88
                _db_update(job_id, stage="translate", progress=prog)

                out = _translate_batch(tr2, batch)
                for src_sent, trg_sent in zip(batch, out):
                    out_parts.append(trg_sent.strip() if trg_sent.strip() else src_sent)
                i += TR_BATCH_SIZE

            return " ".join([p for p in out_parts if p]).strip()

        raise


# ---------------------------
# Export (bilingual)
# ---------------------------
def _export_docx(
    out_path: Path, src_text: str, trg_text: str, src_label: str, trg_label: str
) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Конспект (bilingual)", level=1)

    doc.add_heading(f"Оригинал ({src_label})", level=2)
    doc.add_paragraph(src_text or "")

    doc.add_heading(f"Перевод ({trg_label})", level=2)
    doc.add_paragraph(trg_text or "")

    doc.save(str(out_path))


def _export_pdf(
    out_path: Path, src_text: str, trg_text: str, src_label: str, trg_label: str
) -> None:
    """
    PDF через Platypus + DejaVu — кириллица 100% отображается, переносы нормальные.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    # Важно: шрифт должен существовать в контейнере
    font_regular = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    font_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

    pdfmetrics.registerFont(TTFont("DejaVuSans", font_regular))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", font_bold))

    styles = getSampleStyleSheet()
    base = styles["BodyText"]
    base.fontName = "DejaVuSans"
    base.fontSize = 10
    base.leading = 14

    h1 = styles["Heading1"]
    h1.fontName = "DejaVuSans-Bold"

    h2 = styles["Heading2"]
    h2.fontName = "DejaVuSans-Bold"

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    story = []
    story.append(Paragraph("Конспект (bilingual)", h1))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"Оригинал ({src_label})", h2))
    story.append(Spacer(1, 6))
    story.append(Paragraph((src_text or "").replace("\n", "<br/>"), base))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"Перевод ({trg_label})", h2))
    story.append(Spacer(1, 6))
    story.append(Paragraph((trg_text or "").replace("\n", "<br/>"), base))

    doc.build(story)


# ---------------------------
# Main worker
# ---------------------------
def process_upload_job(job_id: int) -> None:
    base = UPLOADS_ROOT / str(job_id)
    t0 = time.time()

    try:
        _db_update(
            job_id, status="processing", stage="extract", progress=5, error_message=None
        )

        input_path = _find_input_file(base)
        wav_path = base / "audio.wav"

        size, duration = _ffmpeg_extract(input_path, wav_path)
        print(
            f"[uploads_worker] extracted wav in {int(time.time()-t0)}s size={size} duration={duration}",
            flush=True,
        )

        # -------- transcribe --------
        _db_update(job_id, stage="transcribe", progress=30)

        job = _db_get_job(job_id)
        if not job:
            raise RuntimeError("Job not found")

        src_lang_nllb = job.src_language or DEFAULT_SRC_LANG
        _db_update(job_id, src_language=src_lang_nllb)

        stop = _progress_ticker(
            job_id, stage="transcribe", start=31, end=69, step=1, every_sec=3
        )
        try:
            src_text = _transcribe_whisper(wav_path, src_lang_nllb=src_lang_nllb)
        finally:
            stop.set()

        if not src_text.strip():
            raise RuntimeError("ASR produced empty transcript")

        _db_update(job_id, transcript_text=src_text, stage="translate", progress=70)

        # -------- translate --------
        if not job.target_language:
            raise RuntimeError("target_language is not set")

        src_iso = _nllb_to_iso(src_lang_nllb)
        tgt_iso = _nllb_to_iso(job.target_language)
        if not src_iso or not tgt_iso:
            raise RuntimeError(
                f"Unsupported language codes: src={src_lang_nllb}, tgt={job.target_language}"
            )

        translated_text = _translate_text_opus(
            src_text, src_iso=src_iso, tgt_iso=tgt_iso, job_id=job_id
        )
        if not translated_text.strip():
            # лучше так, чем пустой перевод
            translated_text = "[empty translation]"

        _db_update(job_id, translated_text=translated_text, stage="export", progress=90)

        # -------- export --------
        docx_path = base / "result.docx"
        pdf_path = base / "result.pdf"

        _export_docx(
            docx_path,
            src_text,
            translated_text,
            src_label=src_lang_nllb,
            trg_label=job.target_language,
        )
        _export_pdf(
            pdf_path,
            src_text,
            translated_text,
            src_label=src_lang_nllb,
            trg_label=job.target_language,
        )

        _db_update(job_id, status="done", stage="done", progress=100)
        print(
            f"[uploads_worker] done job_id={job_id} src_len={len(src_text)} tr_len={len(translated_text)}",
            flush=True,
        )

    except Exception as e:
        print(f"[uploads_worker] ERROR job_id={job_id}: {e}", flush=True)
        _db_update(
            job_id, status="error", stage="error", progress=100, error_message=str(e)
        )
